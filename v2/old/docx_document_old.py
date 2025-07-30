import io
from docx import Document
from docx.oxml.ns import qn 

import form_checkbox as form_checkbox
import text_replacement as text_replacement

class DocxDocument :
    def __init__(self, bytes) :
        self.docx = None
        if bytes is not None :
            self.docx = Document(io.BytesIO(bytes))

    def get_fields_checkbox(self, includeBody=True, includeHeaders=True, includeFooters=True):
        checkboxes_found = []

        # Namespaces
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'
        }

        # Elementos a buscar
        elements_to_search = []

        if includeBody:
            elements_to_search.append(self.docx._body._element)

        if includeHeaders:
            for section in self.docx.sections:
                if section.header._element is not None:
                    elements_to_search.append(section.header._element)

        if includeFooters:
            for section in self.docx.sections:
                if section.footer._element is not None:
                    elements_to_search.append(section.footer._element)

        # Contador para posiciones únicas por nombre
        name_counters = {}

        for element in elements_to_search:
            # 1. Buscar LEGACY checkboxes: w:fldChar/w:ffData/w:checkBox
            fld_chars = element.findall('.//w:fldChar', namespaces)

            for fld_char in fld_chars:
                ff_data = fld_char.find('w:ffData', namespaces)
                if ff_data is not None and ff_data.find('w:checkBox', namespaces) is not None:
                    # Es un checkbox legacy
                    checkbox_obj = form_checkbox.FormCheckBoxLegacy()

                    # xml_node = todo el contenido de w:fldChar
                    checkbox_obj.xml_node = fld_char

                    # name = valor de w:name/@w:val
                    name_elem = ff_data.find('w:name', namespaces)
                    checkbox_obj.name = name_elem.get(qn('w:val')) if name_elem is not None else ""
                    if checkbox_obj.name and checkbox_obj.name.strip():
                        # Incrementar contador para este nombre
                        if checkbox_obj.name not in name_counters:
                            name_counters[checkbox_obj.name] = 0
                        name_counters[checkbox_obj.name] += 1
                        
                        # xpath = ruta ÚNICA al w:fldChar por nombre y posición
                        position = name_counters[checkbox_obj.name]
                        checkbox_obj.xpath = f"(//w:fldChar[w:ffData/w:name/@w:val='{checkbox_obj.name}' and w:ffData/w:checkBox])[{position}]"

                        # default = estado actual del checkbox
                        checkbox = ff_data.find('w:checkBox', namespaces)
                        default_elem = checkbox.find('w:default', namespaces) if checkbox is not None else None
                        checkbox_obj.default = 1 if (default_elem is not None and default_elem.get(qn('w:val')) == "1") else 0

                        checkboxes_found.append(checkbox_obj)

            # 2. Buscar MODERN checkboxes: w:sdt/w:sdtPr/w14:checkbox
            sdts = element.findall('.//w:sdt', namespaces)

            for sdt in sdts:
                sdt_pr = sdt.find('w:sdtPr', namespaces)
                if sdt_pr is not None and sdt_pr.find('w14:checkbox', namespaces) is not None:
                    # Es un checkbox moderno
                    checkbox_obj = form_checkbox.FormCheckBoxModern()

                    # xml_node = todo el contenido de w:sdt
                    checkbox_obj.xml_node = sdt

                    # tag = valor de w:tag/@w:val
                    tag_elem = sdt_pr.find('w:tag', namespaces)
                    checkbox_obj.tag = tag_elem.get(qn('w:val')) if tag_elem is not None else ""

                    # alias = valor de w:alias/@w:val
                    alias_elem = sdt_pr.find('w:alias', namespaces)
                    checkbox_obj.alias = alias_elem.get(qn('w:val')) if alias_elem is not None else None
                    if (checkbox_obj.tag and checkbox_obj.tag.strip()) or (checkbox_obj.alias and checkbox_obj.alias.strip()):
                        # Usar tag o alias como identificador
                        identifier = checkbox_obj.tag if checkbox_obj.tag and checkbox_obj.tag.strip() else checkbox_obj.alias
                        
                        # Incrementar contador para este identificador
                        if identifier not in name_counters:
                            name_counters[identifier] = 0
                        name_counters[identifier] += 1
                        
                        # xpath = ruta ÚNICA al w:sdt por identificador y posición
                        position = name_counters[identifier]
                        if checkbox_obj.tag and checkbox_obj.tag.strip():
                            checkbox_obj.xpath = f"(//w:sdt[w:sdtPr/w:tag/@w:val='{checkbox_obj.tag}' and w:sdtPr/w14:checkbox])[{position}]"
                        else:
                            checkbox_obj.xpath = f"(//w:sdt[w:sdtPr/w:alias/@w:val='{checkbox_obj.alias}' and w:sdtPr/w14:checkbox])[{position}]"

                        # checked = estado actual del checkbox
                        checkbox_elem = sdt_pr.find('w14:checkbox', namespaces)
                        checked_elem = checkbox_elem.find('w14:checked', namespaces) if checkbox_elem is not None else None
                        checkbox_obj.checked = 1 if (checked_elem is not None and checked_elem.get(qn('w:val')) == "1") else 0

                        checkboxes_found.append(checkbox_obj)

        return checkboxes_found

    def set_field_checkbox_value(self, checkbox_obj, value: bool):
        """
        Activa o desactiva un checkbox modificando directamente el XML del documento
        
        Args:
            checkbox_obj: Objeto FormCheckBoxLegacy o FormCheckBoxModern  
            value: True para activar, False para desactivar
        
        Returns:
            bool: True si se modificó correctamente, False si hubo error
        """
        try:
            # Namespaces
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'
            }
            
            new_val_str = "1" if value else "0"
            
            # Determinar si es legacy o modern
            if hasattr(checkbox_obj, 'name'):  # Legacy
                # Localizar el elemento usando xpath (simulado con find)
                # Buscar por nombre en todo el documento
                body_element = self.docx._body._element
                fld_chars = body_element.findall('.//w:fldChar', namespaces)
                
                found_count = 0
                target_position = int(checkbox_obj.xpath.split('[')[-1].split(']')[0])  # Extraer posición del xpath
                
                for fld_char in fld_chars:
                    ff_data = fld_char.find('w:ffData', namespaces)
                    if ff_data is not None:
                        checkbox = ff_data.find('w:checkBox', namespaces)
                        if checkbox is not None:
                            name_elem = ff_data.find('w:name', namespaces)
                            if name_elem is not None and name_elem.get(qn('w:val')) == checkbox_obj.name:
                                found_count += 1
                                
                                if found_count == target_position:
                                    # Este es el checkbox correcto
                                    default_elem = checkbox.find('w:default', namespaces)
                                    
                                    if default_elem is not None:
                                        default_elem.set(qn('w:val'), new_val_str)
                                    else:
                                        # Crear elemento default si no existe
                                        from docx.oxml import OxmlElement
                                        default_elem = OxmlElement('w:default')
                                        default_elem.set(qn('w:val'), new_val_str)
                                        checkbox.append(default_elem)
                                    
                                    return True
            
            else:  # Modern (tiene tag/alias)
                # Localizar el elemento sdt
                body_element = self.docx._body._element
                sdts = body_element.findall('.//w:sdt', namespaces)
                
                found_count = 0
                identifier = checkbox_obj.tag if checkbox_obj.tag else checkbox_obj.alias
                target_position = int(checkbox_obj.xpath.split('[')[-1].split(']')[0])  # Extraer posición del xpath
                
                for sdt in sdts:
                    sdt_pr = sdt.find('w:sdtPr', namespaces)
                    if sdt_pr is not None:
                        checkbox_elem = sdt_pr.find('w14:checkbox', namespaces)
                        if checkbox_elem is not None:
                            # Verificar tag o alias
                            tag_elem = sdt_pr.find('w:tag', namespaces)
                            alias_elem = sdt_pr.find('w:alias', namespaces)
                            
                            current_identifier = None
                            if tag_elem is not None:
                                current_identifier = tag_elem.get(qn('w:val'))
                            elif alias_elem is not None:
                                current_identifier = alias_elem.get(qn('w:val'))
                            
                            if current_identifier == identifier:
                                found_count += 1
                                
                                if found_count == target_position:
                                    # Este es el checkbox correcto
                                    checked_elem = checkbox_elem.find('w14:checked', namespaces)
                                    
                                    if checked_elem is not None:
                                        checked_elem.set(qn('w:val'), new_val_str)
                                    else:
                                        # Crear elemento checked si no existe
                                        from docx.oxml import OxmlElement
                                        checked_elem = OxmlElement('w14:checked')
                                        checked_elem.set(qn('w:val'), new_val_str)
                                        checkbox_elem.append(checked_elem)
                                    
                                    return True
            
            return False  # No se encontró el checkbox
            
        except Exception as e:
            print(f"Error al modificar checkbox en documento: {e}")
            return False
        
    def get_text_occurrences(self, search_text: str, includeBody=True, includeHeaders=True, includeFooters=True):
        """
        Encuentra todas las ocurrencias de un texto en el documento
        
        Args:
            search_text: Texto a buscar
            includeBody: Buscar en el cuerpo del documento
            includeHeaders: Buscar en headers
            includeFooters: Buscar en footers
        
        Returns:
            List[FormTextReplacement]: Lista de objetos con las ocurrencias encontradas
        """
        occurrences_found = []

        if not search_text:
            return occurrences_found

        # Contador para xpath único
        occurrence_counter = 0

        for section in self.docx.sections:
            if includeBody:
                # Procesar body del documento
                self._find_text_in_part(self.docx, search_text, "body", occurrences_found, occurrence_counter)

            if includeHeaders and section.header._element is not None:
                # Procesar header
                self._find_text_in_part(section.header, search_text, "header", occurrences_found, occurrence_counter)

            if includeFooters and section.footer._element is not None:
                # Procesar footer
                self._find_text_in_part(section.footer, search_text, "footer", occurrences_found, occurrence_counter)

        return occurrences_found

    def _find_text_in_part(self, part, search_text, location, occurrences_found, occurrence_counter):
        """Método auxiliar para buscar texto en una parte del documento"""

        # 1. Procesar párrafos directos
        self._find_text_in_paragraphs(part.paragraphs, search_text, location, "paragraph", occurrences_found, occurrence_counter)

        # 2. Procesar tablas
        self._find_text_in_tables(part.tables, search_text, location, occurrences_found, occurrence_counter)

        # 3. Procesar textboxes
        self._find_text_in_textboxes(part, search_text, location, occurrences_found, occurrence_counter)

    def _find_text_in_paragraphs(self, paragraphs, search_text, location, context, occurrences_found, occurrence_counter):
        """Buscar texto en una lista de párrafos"""
        for paragraph_idx, paragraph in enumerate(paragraphs):
            for run_idx, run in enumerate(paragraph.runs):
                if search_text in run.text:
                    occurrence_counter += 1

                    # Crear objeto FormTextReplacement
                    replacement_obj = text_replacement.FormTextReplacement()
                    replacement_obj.search_text = search_text
                    replacement_obj.replace_text = None  # Se establecerá después
                    replacement_obj.run_node = run
                    replacement_obj.location = location
                    replacement_obj.xpath = f"//{location}/{context}[{paragraph_idx + 1}]/run[{run_idx + 1}]"

                    occurrences_found.append(replacement_obj)

    def _find_text_in_tables(self, tables, search_text, location, occurrences_found, occurrence_counter):
        """Buscar texto en tablas"""
        for table_idx, table in enumerate(tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Procesar párrafos de la celda
                    context = f"table[{table_idx + 1}]/row[{row_idx + 1}]/cell[{cell_idx + 1}]"
                    self._find_text_in_paragraphs(cell.paragraphs, search_text, location, context, occurrences_found, occurrence_counter)

                    # Procesar tablas anidadas
                    self._find_text_in_tables(cell.tables, search_text, location, occurrences_found, occurrence_counter)

    def _find_text_in_textboxes(self, part, search_text, location, occurrences_found, occurrence_counter):
        """Buscar texto en textboxes"""
        try:
            from docx.oxml.ns import qn
            textboxes = part._element.xpath('.//w:txbxContent')

            for textbox_idx, textbox in enumerate(textboxes):
                paragraphs = [paragraph_element for paragraph_element in textbox.iter() if paragraph_element.tag == qn('w:p')]

                for paragraph_idx, paragraph_element in enumerate(paragraphs):
                    try:
                        from docx.text.paragraph import Paragraph
                        paragraph = Paragraph(paragraph_element, part)

                        for run_idx, run in enumerate(paragraph.runs):
                            if search_text in run.text:
                                occurrence_counter += 1

                                replacement_obj = text_replacement.FormTextReplacement()
                                replacement_obj.search_text = search_text
                                replacement_obj.replace_text = None
                                replacement_obj.run_node = run
                                replacement_obj.location = location
                                replacement_obj.xpath = f"//{location}/textbox[{textbox_idx + 1}]/paragraph[{paragraph_idx + 1}]/run[{run_idx + 1}]"

                                occurrences_found.append(replacement_obj)

                    except Exception as e:
                        print(f"Error procesando textbox: {e}")

        except Exception as e:
            print(f"Error buscando en textboxes: {e}")

    def replace_text_occurrence(self, replacement_obj):
        """
        Reemplaza texto en un run específico usando el objeto FormTextReplacement
        
        Args:
            replacement_obj: Objeto FormTextReplacement con run_node, search_text y replace_text
        
        Returns:
            bool: True si se reemplazó correctamente, False si hubo error
        """
        try:
            # Validar que el objeto tenga los datos necesarios
            if replacement_obj.run_node is None:
                print("Error: run_node no está inicializado")
                return False
                
            if not replacement_obj.search_text:
                print("Error: search_text no está definido")
                return False
                
            if replacement_obj.replace_text is None:
                print("Error: replace_text no está definido")
                return False
            
            # Verificar que el texto a buscar aún esté presente en el run
            if replacement_obj.search_text not in replacement_obj.run_node.text:
                print(f"Advertencia: '{replacement_obj.search_text}' no se encontró en el run")
                return False
            
            # Realizar el reemplazo directamente en el run
            # Esto preserva automáticamente todo el formato del run
            original_text = replacement_obj.run_node.text
            replacement_obj.run_node.text = original_text.replace(
                replacement_obj.search_text, 
                replacement_obj.replace_text
            )
            
            return True
            
        except Exception as e:
            print(f"Error al reemplazar texto en run: {e}")
            return False

    def save_to_file(self, file_path):
        """Guarda el documento con las modificaciones"""
        self.docx.save(file_path)

    