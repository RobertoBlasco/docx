"""
Manager para gestionar reemplazo de texto en documentos Word
"""

from docx.oxml.ns import qn
from .base_manager import BaseManager
from models import text_replacement_model as text_replacement


class TextReplacementManager(BaseManager):
    def __init__(self, docx_document):
        super().__init__(docx_document)
    
    def get_text_occurrences(self, search_text: str, includeBody=True, includeHeaders=True, includeFooters=True):
        """
        Encuentra todas las ocurrencias de un texto en el documento
        
        Args:
            search_text: Texto a buscar
            includeBody: Buscar en el cuerpo del documento
            includeHeaders: Buscar en headers
            includeFooters: Buscar en footers
        
        Returns:
            List[FormTextReplacement]: Lista de objetos con las ocurrencias encontradas
        """
        occurrences_found = []
        
        if not search_text:
            return occurrences_found
        
        # Contador para xpath único
        occurrence_counter = 0
        
        for section in self.docx.sections:
            if includeBody:
                # Procesar body del documento
                self._find_text_in_part(self.docx, search_text, "body", occurrences_found, occurrence_counter)
            
            if includeHeaders and section.header._element is not None:
                # Procesar header
                self._find_text_in_part(section.header, search_text, "header", occurrences_found, occurrence_counter)
            
            if includeFooters and section.footer._element is not None:
                # Procesar footer
                self._find_text_in_part(section.footer, search_text, "footer", occurrences_found, occurrence_counter)
        
        return occurrences_found
    
    def _find_text_in_part(self, part, search_text, location, occurrences_found, occurrence_counter):
        """Método auxiliar para buscar texto en una parte del documento"""
        
        # 1. Procesar párrafos directos
        self._find_text_in_paragraphs(part.paragraphs, search_text, location, "paragraph", occurrences_found, occurrence_counter)
        
        # 2. Procesar tablas
        self._find_text_in_tables(part.tables, search_text, location, occurrences_found, occurrence_counter)
        
        # 3. Procesar textboxes
        self._find_text_in_textboxes(part, search_text, location, occurrences_found, occurrence_counter)
    
    def _find_text_in_paragraphs(self, paragraphs, search_text, location, context, occurrences_found, occurrence_counter):
        """Buscar texto en una lista de párrafos"""
        for paragraph_idx, paragraph in enumerate(paragraphs):
            for run_idx, run in enumerate(paragraph.runs):
                if search_text in run.text:
                    occurrence_counter += 1
                    
                    # Crear objeto FormTextReplacement
                    replacement_obj = text_replacement.FormTextReplacement()
                    replacement_obj.search_text = search_text
                    replacement_obj.replace_text = None  # Se establecerá después
                    replacement_obj.run_node = run
                    replacement_obj.location = location
                    replacement_obj.xpath = f"//{location}/{context}[{paragraph_idx + 1}]/run[{run_idx + 1}]"
                    
                    occurrences_found.append(replacement_obj)
    
    def _find_text_in_tables(self, tables, search_text, location, occurrences_found, occurrence_counter):
        """Buscar texto en tablas"""
        for table_idx, table in enumerate(tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Procesar párrafos de la celda
                    context = f"table[{table_idx + 1}]/row[{row_idx + 1}]/cell[{cell_idx + 1}]"
                    self._find_text_in_paragraphs(cell.paragraphs, search_text, location, context, occurrences_found, occurrence_counter)
                    
                    # Procesar tablas anidadas
                    self._find_text_in_tables(cell.tables, search_text, location, occurrences_found, occurrence_counter)
    
    def _find_text_in_textboxes(self, part, search_text, location, occurrences_found, occurrence_counter):
        """Buscar texto en textboxes"""
        try:
            textboxes = part._element.xpath('.//w:txbxContent')
            
            for textbox_idx, textbox in enumerate(textboxes):
                paragraphs = [paragraph_element for paragraph_element in textbox.iter() if paragraph_element.tag == qn('w:p')]
                
                for paragraph_idx, paragraph_element in enumerate(paragraphs):
                    try:
                        from docx.text.paragraph import Paragraph
                        paragraph = Paragraph(paragraph_element, part)
                        
                        for run_idx, run in enumerate(paragraph.runs):
                            if search_text in run.text:
                                occurrence_counter += 1
                                
                                replacement_obj = text_replacement.FormTextReplacement()
                                replacement_obj.search_text = search_text
                                replacement_obj.replace_text = None
                                replacement_obj.run_node = run
                                replacement_obj.location = location
                                replacement_obj.xpath = f"//{location}/textbox[{textbox_idx + 1}]/paragraph[{paragraph_idx + 1}]/run[{run_idx + 1}]"
                                
                                occurrences_found.append(replacement_obj)
                                
                    except Exception as e:
                        print(f"Error procesando textbox: {e}")
                        
        except Exception as e:
            print(f"Error buscando en textboxes: {e}")
    
    def replace_text_occurrence(self, replacement_obj):
        """
        Reemplaza texto en un run específico usando el objeto FormTextReplacement
        
        Args:
            replacement_obj: Objeto FormTextReplacement con run_node, search_text y replace_text
        
        Returns:
            bool: True si se reemplazó correctamente, False si hubo error
        """
        try:
            # Validar que el objeto tenga los datos necesarios
            if replacement_obj.run_node is None:
                print("Error: run_node no está inicializado")
                return False
                
            if not replacement_obj.search_text:
                print("Error: search_text no está definido")
                return False
                
            if replacement_obj.replace_text is None:
                print("Error: replace_text no está definido")
                return False
            
            # Verificar que el texto a buscar aún esté presente en el run
            if replacement_obj.search_text not in replacement_obj.run_node.text:
                print(f"Advertencia: '{replacement_obj.search_text}' no se encontró en el run")
                return False
            
            # Realizar el reemplazo directamente en el run
            # Esto preserva automáticamente todo el formato del run
            original_text = replacement_obj.run_node.text
            replacement_obj.run_node.text = original_text.replace(
                replacement_obj.search_text, 
                replacement_obj.replace_text
            )
            
            return True
            
        except Exception as e:
            print(f"Error al reemplazar texto en run: {e}")
            return False