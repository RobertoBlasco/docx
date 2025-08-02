"""
Manager para gestionar reemplazo de texto por imagen en documentos Word
"""

import io
from docx.shared import Inches
from docx.oxml.ns import qn
from .base_manager import BaseManager
from models import text_to_image_model as text_image_replacement


class TextToImageManager(BaseManager):
    def __init__(self, docx_document):
        super().__init__(docx_document)
    
    def get_text_for_image_replacement(self, search_text: str, includeBody=True, includeHeaders=True, includeFooters=True):
        """
        Encuentra todas las ocurrencias de texto que pueden ser reemplazadas por imagen
        
        Args:
            search_text: Texto a buscar
            includeBody: Buscar en el cuerpo del documento
            includeHeaders: Buscar en headers
            includeFooters: Buscar en footers
        
        Returns:
            List[TextImageReplacement]: Lista de objetos con las ocurrencias encontradas
        """
        replacements_found = []
        
        if not search_text:
            return replacements_found
        
        for section in self.docx.sections:
            if includeBody:
                # Procesar body del documento
                self._find_text_in_part(self.docx, search_text, "body", replacements_found)
            
            if includeHeaders and section.header._element is not None:
                # Procesar header
                self._find_text_in_part(section.header, search_text, "header", replacements_found)
            
            if includeFooters and section.footer._element is not None:
                # Procesar footer
                self._find_text_in_part(section.footer, search_text, "footer", replacements_found)
        
        return replacements_found
    
    def _find_text_in_part(self, part, search_text, location, replacements_found):
        """Método auxiliar para buscar texto en una parte del documento"""
        
        # 1. Procesar párrafos directos
        self._find_text_in_paragraphs(part.paragraphs, search_text, location, "paragraph", replacements_found)
        
        # 2. Procesar tablas
        self._find_text_in_tables(part.tables, search_text, location, replacements_found)
        
        # 3. Procesar textboxes
        self._find_text_in_textboxes(part, search_text, location, replacements_found)
    
    def _find_text_in_paragraphs(self, paragraphs, search_text, location, context, replacements_found):
        """Buscar texto en una lista de párrafos"""
        for paragraph_idx, paragraph in enumerate(paragraphs):
            if search_text in paragraph.text:
                # Crear objeto TextImageReplacement
                replacement_obj = text_image_replacement.TextImageReplacement()
                replacement_obj.search_text = search_text
                replacement_obj.paragraph_node = paragraph
                replacement_obj.location = location
                replacement_obj.xpath = f"//{location}/{context}[{paragraph_idx + 1}]"
                
                replacements_found.append(replacement_obj)
    
    def _find_text_in_tables(self, tables, search_text, location, replacements_found):
        """Buscar texto en tablas"""
        for table_idx, table in enumerate(tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    # Procesar párrafos de la celda
                    context = f"table[{table_idx + 1}]/row[{row_idx + 1}]/cell[{cell_idx + 1}]"
                    self._find_text_in_paragraphs(cell.paragraphs, search_text, location, context, replacements_found)
                    
                    # Procesar tablas anidadas
                    self._find_text_in_tables(cell.tables, search_text, location, replacements_found)
    
    def _find_text_in_textboxes(self, part, search_text, location, replacements_found):
        """Buscar texto en textboxes"""
        try:
            textboxes = part._element.xpath('.//w:txbxContent')
            
            for textbox_idx, textbox in enumerate(textboxes):
                paragraphs = [paragraph_element for paragraph_element in textbox.iter() if paragraph_element.tag == qn('w:p')]
                
                for paragraph_idx, paragraph_element in enumerate(paragraphs):
                    try:
                        from docx.text.paragraph import Paragraph
                        paragraph = Paragraph(paragraph_element, part)
                        
                        if search_text in paragraph.text:
                            replacement_obj = text_image_replacement.TextImageReplacement()
                            replacement_obj.search_text = search_text
                            replacement_obj.paragraph_node = paragraph
                            replacement_obj.location = location
                            replacement_obj.xpath = f"//{location}/textbox[{textbox_idx + 1}]/paragraph[{paragraph_idx + 1}]"
                            
                            replacements_found.append(replacement_obj)
                            
                    except Exception as e:
                        print(f"Error procesando textbox: {e}")
                        
        except Exception as e:
            print(f"Error buscando en textboxes: {e}")
    
    def replace_text_with_image(self, replacement_obj):
        """
        Reemplaza texto por imagen en un párrafo específico
        
        Args:
            replacement_obj: Objeto TextImageReplacement con datos del reemplazo
        
        Returns:
            bool: True si se reemplazó correctamente, False si hubo error
        """
        try:
            # Validar que el objeto tenga los datos necesarios
            if replacement_obj.paragraph_node is None:
                print("Error: paragraph_node no está inicializado")
                return False
                
            if not replacement_obj.search_text:
                print("Error: search_text no está definido")
                return False
                
            if replacement_obj.image_data is None:
                print("Error: image_data no está definido")
                return False
            
            paragraph = replacement_obj.paragraph_node
            search_text = replacement_obj.search_text
            full_text = paragraph.text
            
            # Verificar que el texto aún esté presente
            if search_text not in full_text:
                print(f"Advertencia: '{search_text}' no se encontró en el párrafo")
                return False
            
            # Encontrar todas las posiciones donde aparece el texto
            positions = []
            start = 0
            while True:
                pos = full_text.find(search_text, start)
                if pos == -1:
                    break
                positions.append(pos)
                start = pos + len(search_text)
            
            if not positions:
                return False
            
            # Reconstruir el párrafo con las imágenes
            self._rebuild_paragraph_with_images(paragraph, full_text, positions, replacement_obj)
            
            return True
            
        except Exception as e:
            print(f"Error al reemplazar texto con imagen: {e}")
            return False
    
    def _rebuild_paragraph_with_images(self, paragraph, full_text: str, positions: list, replacement_obj):
        """Reconstruye el párrafo reemplazando texto por imágenes"""
        search_text = replacement_obj.search_text
        image_data = replacement_obj.image_data
        width = replacement_obj.width
        height = replacement_obj.height
        
        # Crear una lista de elementos para reconstruir
        elements = []
        current_pos = 0
        
        for pos in positions:
            # Texto antes del marcador
            before_text = full_text[current_pos:pos]
            if before_text:
                elements.append({"type": "text", "content": before_text})
            
            # Imagen en lugar del texto
            elements.append({
                "type": "image",
                "data": image_data,
                "width": width,
                "height": height
            })
            
            current_pos = pos + len(search_text)
        
        # Texto después del último marcador
        after_text = full_text[current_pos:]
        if after_text:
            elements.append({"type": "text", "content": after_text})
        
        # Limpiar runs existentes preservando el formato del párrafo
        self._clear_paragraph_runs(paragraph)
        
        # Reconstruir el párrafo
        for element in elements:
            if element["type"] == "text":
                if element["content"]:  # Solo agregar si no está vacío
                    paragraph.add_run(element["content"])
            elif element["type"] == "image":
                self._add_image_to_paragraph(paragraph, element["data"], element["width"], element["height"])
    
    def _clear_paragraph_runs(self, paragraph):
        """Limpia todos los runs de un párrafo preservando el formato"""
        # Eliminar todos los runs existentes
        for run in paragraph.runs[::-1]:  # Iterar en reverso para evitar problemas de índices
            paragraph._element.remove(run._element)
    
    def _add_image_to_paragraph(self, paragraph, image_data: bytes, width: float, height: float):
        """Agrega una imagen a un párrafo"""
        try:
            # Crear un nuevo run para la imagen
            run = paragraph.add_run()
            
            # Convertir dimensiones a pulgadas
            width_inches = self._pixels_to_inches(width) if width else None
            height_inches = self._pixels_to_inches(height) if height else None
            
            # Crear stream de la imagen
            image_stream = io.BytesIO(image_data)
            
            # Agregar la imagen al run
            run.add_picture(
                image_stream,
                width=Inches(width_inches) if width_inches else None,
                height=Inches(height_inches) if height_inches else None
            )
            
        except Exception as e:
            print(f"Error agregando imagen al párrafo: {e}")
    
    def _pixels_to_inches(self, pixels: float) -> float:
        """Convierte píxeles a pulgadas (asumiendo 96 DPI)"""
        if pixels is None:
            return None
        return pixels / 96.0