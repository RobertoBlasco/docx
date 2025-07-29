import os
import time

from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def replace_text_with_text(doc: Document, action) -> float:
    start_time = time.time()
    for section in doc.sections:
         process_part(doc, action)               # Body del documento
         process_part(section.header, action)    # Headers
         process_part(section.footer, action)    # Footers
    end_time = time.time()
    elapsed = end_time - start_time
    num_seconds = float(f"{elapsed:.4f}")
    return num_seconds

def process_part(part, action):
    process_paragraphs(part.paragraphs, action) # Párrafos
    process_tables(part.tables, action)         # Tablas
    process_textboxes(part, action)             # TextBoxes
    
def process_paragraphs(paragraphs, action):
    for paragraph in paragraphs:
        replace_in_paragraph(paragraph, action)

def process_tables(tables, action):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs, action)
                process_tables(cell.tables, action)

def process_textboxes(part, action):
    textboxes = part._element.xpath('.//w:txbxContent')
    for textbox in textboxes:
        paragraphs = [paragraph_element for paragraph_element in textbox.iter() if paragraph_element.tag == qn('w:p')]
        for paragraph_element in paragraphs:
            try:
                paragraph = create_paragraph_from_xml(paragraph_element, part)
                replace_in_paragraph(paragraph, action)
            except Exception as e:
                print(f"Error procesando textbox: {e}")
                
def create_paragraph_from_xml(xml_element, part):
    from docx.text.paragraph import Paragraph
    return Paragraph(xml_element, part)

def replace_in_paragraph(paragraph, action):
    search_text = action.search_text    # Texto a buscar
    replace_text = action.replace_text  # Texto de reemplazo
    if search_text and replace_text :
    
        # Recorrer cada run (fragmento de texto con formato uniforme) del párrafo
        for run in paragraph.runs:
            # Verificar si el texto a buscar está presente en este run
            if search_text in run.text:
                # Reemplazar directamente en el texto del run
                # Esto preserva automáticamente todo el formato original del run
                # (negrita, cursiva, subrayado, fuente, tamaño, color, etc.)
                run.text = run.text.replace(search_text, replace_text)

        # # Recorrer todos los runs del párrafo
        # for run in paragraph.runs:
        #     if label in run.text:
        #         # Acceder al elemento XML directamente
        #         xml_element = run._element

        #         # Buscar el primer nodo <w:t>
        #         t_node = xml_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
        #         if t_node is not None and label in t_node.text:
        #             t_node.text = t_node.text.replace(label, text)