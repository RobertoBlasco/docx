#!/usr/bin/env python3
"""
ActionReplaceTextWithImage - Implementación usando solo python-docx
Reemplaza la implementación anterior que usaba xpath y manipulaciones XML externas
"""

import os
import io
import logging
from action_docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph#!/usr/bin/env python3
"""
ActionReplaceTextWithImage - Implementación usando solo python-docx
Reemplaza la implementación anterior que usaba xpath y manipulaciones XML externas
"""

import os
import io
import logging
from action_docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.text.run import Run

logger = logging.getLogger("IneoDocx")

class ActionReplaceTextWithImage:
    """
    Acción para reemplazar texto con imágenes usando solo python-docx
    """
    
    def __init__(self, document: Document):
        """
        Inicializa la acción con el documento a modificar
        
        Args:
            document: Documento python-docx ya cargado
        """
        self.document = document
        
    def replace_text_with_image(self, search_text: str, image_data: bytes, width: float = None, height: float = None):
        """
        Reemplaza todas las ocurrencias de un texto con una imagen
        
        Args:
            search_text: Texto a buscar y reemplazar
            image_data: Datos binarios de la imagen
            width: Ancho en píxeles (opcional)
            height: Alto en píxeles (opcional)
            
        Returns:
            int: Número de reemplazos realizados
        """
        if not image_data:
            logger.error("No se proporcionaron datos de imagen")
            return 0
            
        replacements = 0
        
        logger.info(f"Buscando texto '{search_text}' para reemplazar con imagen")
        
        # Procesar párrafos del documento principal
        replacements += self._process_paragraphs(self.document.paragraphs, search_text, image_data, width, height)
        
        # Procesar tablas
        replacements += self._process_tables(self.document.tables, search_text, image_data, width, height)
        
        # Procesar headers y footers de todas las secciones
        for section in self.document.sections:
            # Header
            if section.header:
                replacements += self._process_paragraphs(section.header.paragraphs, search_text, image_data, width, height)
                replacements += self._process_tables(section.header.tables, search_text, image_data, width, height)
            
            # Footer
            if section.footer:
                replacements += self._process_paragraphs(section.footer.paragraphs, search_text, image_data, width, height)
                replacements += self._process_tables(section.footer.tables, search_text, image_data, width, height)
        
        logger.info(f"Realizados {replacements} reemplazos de texto con imagen")
        return replacements
    
    def _process_paragraphs(self, paragraphs, search_text: str, image_data: bytes, width: float, height: float):
        """
        Procesa una lista de párrafos buscando texto a reemplazar
        
        Args:
            paragraphs: Lista de párrafos a procesar
            search_text: Texto a buscar
            image_data: Datos de la imagen
            width: Ancho en píxeles
            height: Alto en píxeles
            
        Returns:
            int: Número de reemplazos realizados
        """
        replacements = 0
        
        for paragraph in paragraphs:
            if search_text in paragraph.text:
                replacements += self._replace_in_paragraph(paragraph, search_text, image_data, width, height)
        
        return replacements
    
    def _process_tables(self, tables, search_text: str, image_data: bytes, width: float, height: float):
        """
        Procesa todas las tablas del documento
        
        Args:
            tables: Lista de tablas a procesar
            search_text: Texto a buscar
            image_data: Datos de la imagen
            width: Ancho en píxeles
            height: Alto en píxeles
            
        Returns:
            int: Número de reemplazos realizados
        """
        replacements = 0
        
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    # Procesar párrafos en la celda
                    replacements += self._process_paragraphs(cell.paragraphs, search_text, image_data, width, height)
                    # Procesar tablas anidadas
                    replacements += self._process_tables(cell.tables, search_text, image_data, width, height)
        
        return replacements
    
    def _replace_in_paragraph(self, paragraph, search_text: str, image_data: bytes, width: float, height: float):
        """
        Reemplaza texto por imagen en un párrafo específico
        
        Args:
            paragraph: Párrafo a procesar
            search_text: Texto a buscar
            image_data: Datos de la imagen
            width: Ancho en píxeles
            height: Alto en píxeles
            
        Returns:
            int: Número de reemplazos realizados
        """
        full_text = paragraph.text
        
        if search_text not in full_text:
            return 0
        
        # Encontrar todas las posiciones donde aparece el texto
        positions = []
        start = 0
        while True:
            pos = full_text.find(search_text, start)
            if pos == -1:
                break
            positions.append(pos)
            start = pos + len(search_text)
        
        if not positions:
            return 0
        
        # Reconstruir el párrafo con las imágenes
        self._rebuild_paragraph_with_images(paragraph, full_text, search_text, positions, image_data, width, height)
        
        return len(positions)
    
    def _rebuild_paragraph_with_images(self, paragraph, full_text: str, search_text: str, positions: list, image_data: bytes, width: float, height: float):
        """
        Reconstruye un párrafo reemplazando texto con imágenes
        
        Args:
            paragraph: Párrafo a reconstruir
            full_text: Texto completo del párrafo
            search_text: Texto a reemplazar
            positions: Lista de posiciones donde se encontró el texto
            image_data: Datos de la imagen
            width: Ancho en píxeles
            height: Alto en píxeles
        """
        # Crear una lista de elementos para reconstruir
        elements = []
        current_pos = 0
        
        for pos in positions:
            # Texto antes del marcador
            before_text = full_text[current_pos:pos]
            if before_text:
                elements.append({"type": "text", "content": before_text})
            
            # Imagen en lugar del texto
            elements.append({
                "type": "image",
                "data": image_data,
                "width": width,
                "height": height
            })
            
            current_pos = pos + len(search_text)
        
        # Texto después del último marcador
        after_text = full_text[current_pos:]
        if after_text:
            elements.append({"type": "text", "content": after_text})
        
        # Limpiar runs existentes preservando el formato del párrafo
        self._clear_paragraph_runs(paragraph)
        
        # Reconstruir el párrafo
        for element in elements:
            if element["type"] == "text":
                if element["content"]:  # Solo agregar si no está vacío
                    paragraph.add_run(element["content"])
            elif element["type"] == "image":
                self._add_image_to_paragraph(paragraph, element["data"], element["width"], element["height"])
    
    def _clear_paragraph_runs(self, paragraph):
        """
        Limpia todos los runs de un párrafo preservando la estructura
        
        Args:
            paragraph: Párrafo a limpiar
        """
        # Limpiar el texto de todos los runs existentes
        for run in paragraph.runs:
            run.text = ""
        
        # Alternativamente, podríamos remover los runs completamente
        # Pero esto preserva mejor el formato
    
    def _add_image_to_paragraph(self, paragraph, image_data: bytes, width: float, height: float):
        """
        Agrega una imagen a un párrafo
        
        Args:
            paragraph: Párrafo donde agregar la imagen
            image_data: Datos binarios de la imagen
            width: Ancho en píxeles
            height: Alto en píxeles
        """
        try:
            # Crear un nuevo run para la imagen
            run = paragraph.add_run()
            
            # Convertir dimensiones a pulgadas
            width_inches = self._pixels_to_inches(width) if width else None
            height_inches = self._pixels_to_inches(height) if height else None
            
            # Crear stream de la imagen
            image_stream = io.BytesIO(image_data)
            
            # Agregar la imagen al run
            run.add_picture(
                image_stream,
                width=Inches(width_inches) if width_inches else None,
                height=Inches(height_inches) if height_inches else None
            )
            
            logger.info(f"Imagen agregada al párrafo: {width_inches}x{height_inches} inches")
            
        except Exception as e:
            logger.error(f"Error agregando imagen al párrafo: {e}")
    
    def _pixels_to_inches(self, pixels: float, dpi: int = 96) -> float:
        """
        Convierte píxeles a pulgadas
        
        Args:
            pixels: Valor en píxeles
            dpi: Puntos por pulgada (por defecto 96)
            
        Returns:
            float: Valor en pulgadas
        """
        if pixels is None:
            return None
        
        try:
            return float(pixels) / dpi
        except (ValueError, TypeError):
            logger.warning(f"No se pudo convertir {pixels} píxeles a pulgadas")
            return None

def replace_text_with_image(doc: Document, action, xml_data):
    """
    Función standalone para compatibilidad con la implementación anterior
    
    Args:
        doc: Documento python-docx
        action: Acción ActionReplaceTextWithImage
        xml_data: Datos XML con las imágenes disponibles
        
    Returns:
        bool: True si se realizó al menos un reemplazo
    """
    # Crear instancia de la clase
    replacer = ActionReplaceTextWithImage(doc)
    
    # Obtener datos de la imagen
    image_data = get_image_data(xml_data, action.image_id)
    if not image_data:
        logger.error(f"No se encontró imagen con ID: {action.image_id}")
        return False
    
    # Realizar reemplazos
    replacements = replacer.replace_text_with_image(
        search_text=action.search_text,
        image_data=image_data,
        width=action.width,
        height=action.height
    )
    
    return replacements > 0

def get_image_data(xml_data, image_id):
    """
    Obtiene los datos de una imagen por su ID desde dataStorage
    
    Args:
        xml_data: Datos XML con las imágenes procesadas
        image_id: ID de la imagen a buscar
        
    Returns:
        bytes: Datos de la imagen o None si no se encuentra
    """
    # Buscar la imagen en xml_data.images
    target_image = None
    for img in xml_data.images:
        if img.id == image_id:
            target_image = img
            break
    
    if not target_image:
        logger.error(f"Imagen con ID {image_id} no encontrada")
        return None
    
    try:
        # Después de process_xml_data(), todas las imágenes están en dataStorage
        # con nombre MD5 y path actualizado a FILE://dataStorage/MD5
        if target_image.path.startswith("FILE://"):
            file_path = target_image.path[7:]  # Remover FILE://
            
            # Verificar que el archivo existe en dataStorage
            if not os.path.exists(file_path):
                logger.error(f"Archivo de imagen no existe en dataStorage: {file_path}")
                return None
            
            # Cargar imagen desde dataStorage
            with open(file_path, 'rb') as f:
                image_data = f.read()
                logger.info(f"Imagen cargada desde dataStorage: {file_path} ({len(image_data)} bytes)")
                return image_data
                
        else:
            # Esto no debería ocurrir después de process_xml_data()
            logger.error(f"Imagen no procesada correctamente por dataStorage: {target_image.path}")
            return None
            
    except Exception as e:
        logger.error(f"Error cargando imagen {image_id} desde dataStorage: {e}")
        return None
from docx.text.run import Run

logger = logging.getLogger("IneoDocx")

class ActionReplaceTextWithImage:
    """
    Acción para reemplazar texto con imágenes usando solo python-docx
    """
    
    def __init__(self, document: Document):
        """
        Inicializa la acción con el documento a modificar
        
        Args:
            document: Documento python-docx ya cargado
        """
        self.document = document
        
    def replace_text_with_image(self, search_text: str, image_data: bytes, width: float = None, height: float = None):
        """
        Reemplaza todas las ocurrencias de un texto con una imagen
        
        Args:
            search_text: Texto a buscar y reemplazar
            image_data: Datos binarios de la imagen
            width: Ancho en píxeles (opcional)
            height: Alto en píxeles (opcional)
            
        Returns:
            int: Número de reemplazos realizados
        """
        if not image_data:
            logger.error("No se proporcionaron datos de imagen")
            return 0
            
        replacements = 0
        
        logger.info(f"Buscando texto '{search_text}' para reemplazar con imagen")
        
        # Procesar párrafos del documento principal
        replacements += self._process_paragraphs(self.document.paragraphs, search_text, image_data, width, height)
        
        # Procesar tablas
        replacements += self._process_tables(self.document.tables, search_text, image_data, width, height)
        
        # Procesar headers y footers de todas las secciones
        for section in self.document.sections:
            # Header
            if section.header:
                replacements += self._process_paragraphs(section.header.paragraphs, search_text, image_data, width, height)
                replacements += self._process_tables(section.header.tables, search_text, image_data, width, height)
            
            # Footer
            if section.footer:
                replacements += self._process_paragraphs(section.footer.paragraphs, search_text, image_data, width, height)
                replacements += self._process_tables(section.footer.tables, search_text, image_data, width, height)
        
        logger.info(f"Realizados {replacements} reemplazos de texto con imagen")
        return replacements
    
    def _process_paragraphs(self, paragraphs, search_text: str, image_data: bytes, width: float, height: float):
        """
        Procesa una lista de párrafos buscando texto a reemplazar
        
        Args:
            paragraphs: Lista de párrafos a procesar
            search_text: Texto a buscar
            image_data: Datos de la imagen
            width: Ancho en píxeles
            height: Alto en píxeles
            
        Returns:
            int: Número de reemplazos realizados
        """
        replacements = 0
        
        for paragraph in paragraphs:
            if search_text in paragraph.text:
                replacements += self._replace_in_paragraph(paragraph, search_text, image_data, width, height)
        
        return replacements
    
    def _process_tables(self, tables, search_text: str, image_data: bytes, width: float, height: float):
        """
        Procesa todas las tablas del documento
        
        Args:
            tables: Lista de tablas a procesar
            search_text: Texto a buscar
            image_data: Datos de la imagen
            width: Ancho en píxeles
            height: Alto en píxeles
            
        Returns:
            int: Número de reemplazos realizados
        """
        replacements = 0
        
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    # Procesar párrafos en la celda
                    replacements += self._process_paragraphs(cell.paragraphs, search_text, image_data, width, height)
                    # Procesar tablas anidadas
                    replacements += self._process_tables(cell.tables, search_text, image_data, width, height)
        
        return replacements
    
    def _replace_in_paragraph(self, paragraph, search_text: str, image_data: bytes, width: float, height: float):
        """
        Reemplaza texto por imagen en un párrafo específico
        
        Args:
            paragraph: Párrafo a procesar
            search_text: Texto a buscar
            image_data: Datos de la imagen
            width: Ancho en píxeles
            height: Alto en píxeles
            
        Returns:
            int: Número de reemplazos realizados
        """
        full_text = paragraph.text
        
        if search_text not in full_text:
            return 0
        
        # Encontrar todas las posiciones donde aparece el texto
        positions = []
        start = 0
        while True:
            pos = full_text.find(search_text, start)
            if pos == -1:
                break
            positions.append(pos)
            start = pos + len(search_text)
        
        if not positions:
            return 0
        
        # Reconstruir el párrafo con las imágenes
        self._rebuild_paragraph_with_images(paragraph, full_text, search_text, positions, image_data, width, height)
        
        return len(positions)
    
    def _rebuild_paragraph_with_images(self, paragraph, full_text: str, search_text: str, positions: list, image_data: bytes, width: float, height: float):
        """
        Reconstruye un párrafo reemplazando texto con imágenes
        
        Args:
            paragraph: Párrafo a reconstruir
            full_text: Texto completo del párrafo
            search_text: Texto a reemplazar
            positions: Lista de posiciones donde se encontró el texto
            image_data: Datos de la imagen
            width: Ancho en píxeles
            height: Alto en píxeles
        """
        # Crear una lista de elementos para reconstruir
        elements = []
        current_pos = 0
        
        for pos in positions:
            # Texto antes del marcador
            before_text = full_text[current_pos:pos]
            if before_text:
                elements.append({"type": "text", "content": before_text})
            
            # Imagen en lugar del texto
            elements.append({
                "type": "image",
                "data": image_data,
                "width": width,
                "height": height
            })
            
            current_pos = pos + len(search_text)
        
        # Texto después del último marcador
        after_text = full_text[current_pos:]
        if after_text:
            elements.append({"type": "text", "content": after_text})
        
        # Limpiar runs existentes preservando el formato del párrafo
        self._clear_paragraph_runs(paragraph)
        
        # Reconstruir el párrafo
        for element in elements:
            if element["type"] == "text":
                if element["content"]:  # Solo agregar si no está vacío
                    paragraph.add_run(element["content"])
            elif element["type"] == "image":
                self._add_image_to_paragraph(paragraph, element["data"], element["width"], element["height"])
    
    def _clear_paragraph_runs(self, paragraph):
        """
        Limpia todos los runs de un párrafo preservando la estructura
        
        Args:
            paragraph: Párrafo a limpiar
        """
        # Limpiar el texto de todos los runs existentes
        for run in paragraph.runs:
            run.text = ""
        
        # Alternativamente, podríamos remover los runs completamente
        # Pero esto preserva mejor el formato
    
    def _add_image_to_paragraph(self, paragraph, image_data: bytes, width: float, height: float):
        """
        Agrega una imagen a un párrafo
        
        Args:
            paragraph: Párrafo donde agregar la imagen
            image_data: Datos binarios de la imagen
            width: Ancho en píxeles
            height: Alto en píxeles
        """
        try:
            # Crear un nuevo run para la imagen
            run = paragraph.add_run()
            
            # Convertir dimensiones a pulgadas
            width_inches = self._pixels_to_inches(width) if width else None
            height_inches = self._pixels_to_inches(height) if height else None
            
            # Crear stream de la imagen
            image_stream = io.BytesIO(image_data)
            
            # Agregar la imagen al run
            run.add_picture(
                image_stream,
                width=Inches(width_inches) if width_inches else None,
                height=Inches(height_inches) if height_inches else None
            )
            
            logger.info(f"Imagen agregada al párrafo: {width_inches}x{height_inches} inches")
            
        except Exception as e:
            logger.error(f"Error agregando imagen al párrafo: {e}")
    
    def _pixels_to_inches(self, pixels: float, dpi: int = 96) -> float:
        """
        Convierte píxeles a pulgadas
        
        Args:
            pixels: Valor en píxeles
            dpi: Puntos por pulgada (por defecto 96)
            
        Returns:
            float: Valor en pulgadas
        """
        if pixels is None:
            return None
        
        try:
            return float(pixels) / dpi
        except (ValueError, TypeError):
            logger.warning(f"No se pudo convertir {pixels} píxeles a pulgadas")
            return None

def replace_text_with_image(doc: Document, action, xml_data):
    """
    Función standalone para compatibilidad con la implementación anterior
    
    Args:
        doc: Documento python-docx
        action: Acción ActionReplaceTextWithImage
        xml_data: Datos XML con las imágenes disponibles
        
    Returns:
        bool: True si se realizó al menos un reemplazo
    """
    # Crear instancia de la clase
    replacer = ActionReplaceTextWithImage(doc)
    
    # Obtener datos de la imagen
    image_data = get_image_data(xml_data, action.image_id)
    if not image_data:
        logger.error(f"No se encontró imagen con ID: {action.image_id}")
        return False
    
    # Realizar reemplazos
    replacements = replacer.replace_text_with_image(
        search_text=action.search_text,
        image_data=image_data,
        width=action.width,
        height=action.height
    )
    
    return replacements > 0

def get_image_data(xml_data, image_id):
    """
    Obtiene los datos de una imagen por su ID desde dataStorage
    
    Args:
        xml_data: Datos XML con las imágenes procesadas
        image_id: ID de la imagen a buscar
        
    Returns:
        bytes: Datos de la imagen o None si no se encuentra
    """
    # Buscar la imagen en xml_data.images
    target_image = None
    for img in xml_data.images:
        if img.id == image_id:
            target_image = img
            break
    
    if not target_image:
        logger.error(f"Imagen con ID {image_id} no encontrada")
        return None
    
    try:
        # Después de process_xml_data(), todas las imágenes están en dataStorage
        # con nombre MD5 y path actualizado a FILE://dataStorage/MD5
        if target_image.path.startswith("FILE://"):
            file_path = target_image.path[7:]  # Remover FILE://
            
            # Verificar que el archivo existe en dataStorage
            if not os.path.exists(file_path):
                logger.error(f"Archivo de imagen no existe en dataStorage: {file_path}")
                return None
            
            # Cargar imagen desde dataStorage
            with open(file_path, 'rb') as f:
                image_data = f.read()
                logger.info(f"Imagen cargada desde dataStorage: {file_path} ({len(image_data)} bytes)")
                return image_data
                
        else:
            # Esto no debería ocurrir después de process_xml_data()
            logger.error(f"Imagen no procesada correctamente por dataStorage: {target_image.path}")
            return None
            
    except Exception as e:
        logger.error(f"Error cargando imagen {image_id} desde dataStorage: {e}")
        return None