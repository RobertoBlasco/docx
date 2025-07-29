import io
import time

from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def replace_text_with_image(doc: Document, action) -> float:
    start_time = time.time()
    for section in doc.sections:
         process_part(doc, action)               # Body del documento
         process_part(section.header, action)    # Headers
         process_part(section.footer, action)    # Footers
    end_time = time.time()
    elapsed = end_time - start_time
    num_seconds = float(f"{elapsed:.4f}")
    return num_seconds

def process_part(part, action):
    process_paragraphs(part.paragraphs, action) # Párrafos
    process_tables(part.tables, action)         # Tablas
    process_textboxes(part, action)             # TextBoxes
    
def process_paragraphs(paragraphs, action):
    for paragraph in paragraphs:
        replace_in_paragraph(paragraph, action)

def process_tables(tables, action):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs, action)
                process_tables(cell.tables, action)

def process_textboxes(part, action):
    textboxes = part._element.xpath('.//w:txbxContent')
    for textbox in textboxes:
        paragraphs = [paragraph_element for paragraph_element in textbox.iter() if paragraph_element.tag == qn('w:p')]
        for paragraph_element in paragraphs:
            try:
                paragraph = create_paragraph_from_xml(paragraph_element, part)
                replace_in_paragraph(paragraph, action)
            except Exception as e:
                print(f"Error procesando textbox: {e}")
                
def create_paragraph_from_xml(xml_element, part):
    from docx.text.paragraph import Paragraph
    return Paragraph(xml_element, part)


def replace_in_paragraph(paragraph, action):
        full_text = paragraph.text
        search_text = action.search_text
        if search_text not in full_text:
            return 0
        
        # Encontrar todas las posiciones donde aparece el texto
        positions = []
        start = 0
        while True:
            pos = full_text.find(search_text, start)
            if pos == -1:
                break
            positions.append(pos)
            start = pos + len(search_text)
        
        if not positions:
            return 0
        
        # Reconstruir el párrafo con las imágenes
        rebuild_paragraph_with_images(paragraph, full_text, positions, action)
        
        return len(positions)
    
def rebuild_paragraph_with_images(self, paragraph, full_text: str, positions: list, action):
        search_text = action.search_text
        image_data = action.image_data
        width = action.width
        height = action.height
        
        # Crear una lista de elementos para reconstruir
        elements = []
        current_pos = 0
        
        for pos in positions:
            # Texto antes del marcador
            before_text = full_text[current_pos:pos]
            if before_text:
                elements.append({"type": "text", "content": before_text})
            
            # Imagen en lugar del texto
            elements.append({
                "type": "image",
                "data": image_data,
                "width": width,
                "height": height
            })
            
            current_pos = pos + len(search_text)
        
        # Texto después del último marcador
        after_text = full_text[current_pos:]
        if after_text:
            elements.append({"type": "text", "content": after_text})
        
        # Limpiar runs existentes preservando el formato del párrafo
        self._clear_paragraph_runs(paragraph)
        
        # Reconstruir el párrafo
        for element in elements:
            if element["type"] == "text":
                if element["content"]:  # Solo agregar si no está vacío
                    paragraph.add_run(element["content"])
            elif element["type"] == "image":
                add_image_to_paragraph(paragraph, element["data"], element["width"], element["height"])
                
def add_image_to_paragraph(self, paragraph, image_data: bytes, width: float, height: float):
        try:
            # Crear un nuevo run para la imagen
            run = paragraph.add_run()
            
            # Convertir dimensiones a pulgadas
            width_inches = self._pixels_to_inches(width) if width else None
            height_inches = self._pixels_to_inches(height) if height else None
            
            # Crear stream de la imagen
            image_stream = io.BytesIO(image_data)
            
            # Agregar la imagen al run
            run.add_picture(
                image_stream,
                width=Inches(width_inches) if width_inches else None,
                height=Inches(height_inches) if height_inches else None
            )
            
            #logger.info(f"Imagen agregada al párrafo: {width_inches}x{height_inches} inches")
            
        except Exception as e:
            #logger.error(f"Error agregando imagen al párrafo: {e}")
            pass