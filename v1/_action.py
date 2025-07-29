#!/usr/bin/env python3
"""
ActionReplaceTextWithText
"""

import logging
from action_docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger("IneoDocx")

class ActionReplaceTextWithText:
    """
    Acción para reemplazar texto en un documento usando python-docx
    """

    def __init__(self, document_path: str):
        """
        Inicializa la acción con el documento a modificar
        
        Args:
            document_path: Ruta al documento .docx
        """
        self.document_path = document_path
        self.document = None

    def load_document(self):
        """Carga el documento Word"""
        try:
            self.document = Document(self.document_path)
            logger.info(f"Documento cargado: {self.document_path}")
            return True
        except Exception as e:
            logger.error(f"Error cargando documento: {e}")
            return False
        
    def replace_text_with_text(doc, action):
    
        for section in doc.sections:
            process_part(doc, action)               # Body del documento
            process_part(section.header, action)    # Headers
            process_part(section.footer, action)    # Footers

def process_part(part, action):
    process_paragraphs(part.paragraphs, action) # Párrafos
    process_tables(part.tables, action)         # Tablas
    process_textboxes(part, action)             # TextBoxes
    
def process_paragraphs(paragraphs, action):
    for paragraph in paragraphs:
        replace_in_paragraph(paragraph, action)

def process_tables(tables, action):
    label = action.label
    text = action.text
    
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs, action)
                process_tables(cell.tables, action)

def process_textboxes(part, action):
    label = action.label
    text = action.text
    
    textboxes = part._element.xpath('.//w:txbxContent')
    for textbox in textboxes:
        paragraphs = [paragraph_element for paragraph_element in textbox.iter() if paragraph_element.tag == qn('w:p')]
        for paragraph_element in paragraphs:
            try:
                paragraph = create_paragraph_from_xml(paragraph_element, part)
                replace_in_paragraph(paragraph, action)
            except Exception as e:
                print(f"Error procesando textbox: {e}")
                
def create_paragraph_from_xml(xml_element, part):
    
    from docx.text.paragraph import Paragraph
    return Paragraph(xml_element, part)

def replace_in_paragraph(paragraph, action):
    """
    Reemplaza texto en un párrafo preservando el formato original.
    
    Esta función utiliza un enfoque directo que mantiene automáticamente
    el formato (negrita, cursiva, fuente, color, etc.) al reemplazar
    texto directamente en cada run individual.
    
    Args:
        paragraph: Párrafo del documento donde realizar el reemplazo
        action: Acción que contiene el texto a buscar (label) y el texto de reemplazo (text)
    """
    label = action.label  # Texto a buscar
    text = action.text    # Texto de reemplazo
    
    # Recorrer cada run (fragmento de texto con formato uniforme) del párrafo
    for run in paragraph.runs:
        # Verificar si el texto a buscar está presente en este run
        if label in run.text:
            # Reemplazar directamente en el texto del run
            # Esto preserva automáticamente todo el formato original del run
            # (negrita, cursiva, subrayado, fuente, tamaño, color, etc.)
            run.text = run.text.replace(label, text)

    # # Recorrer todos los runs del párrafo
    # for run in paragraph.runs:
    #     if label in run.text:
    #         # Acceder al elemento XML directamente
    #         xml_element = run._element

    #         # Buscar el primer nodo <w:t>
    #         t_node = xml_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
    #         if t_node is not None and label in t_node.text:
    #             t_node.text = t_node.text.replace(label, text)